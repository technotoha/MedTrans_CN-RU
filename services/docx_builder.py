"""
DOCX document builder service.
Reconstructs translated content into a properly formatted Word document.
"""

import asyncio
from pathlib import Path
from typing import List, Dict, Optional, Any
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from models.document_models import (
    PageContent, DocumentElement, ElementType, TableData, ProcessingStatus
)
from utils.logging_utils import get_logger, timer


logger = get_logger(__name__)


class DOCXBuilder:
    """
    Builds DOCX documents from translated content.
    
    Preserves document structure including:
    - Headers and paragraphs
    - Tables with formatting
    - Images (with OCR captions)
    - Lists
    """
    
    def __init__(self):
        """Initialize the DOCX builder."""
        self.logger = logger
        self.doc = None
        
        # Style configurations
        self.styles = {
            "header1": {"size": 18, "bold": True, "space_after": 12},
            "header2": {"size": 14, "bold": True, "space_after": 10},
            "header3": {"size": 12, "bold": True, "space_after": 8},
            "normal": {"size": 11, "bold": False, "space_after": 6},
            "caption": {"size": 9, "italic": True, "space_after": 6},
            "table_text": {"size": 10, "bold": False}
        }
    
    @timer(logger, "DOCX building")
    async def build_document(
        self,
        pages: List[PageContent],
        output_path: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Path:
        """
        Build a DOCX document from translated pages.
        
        Args:
            pages: List of page contents with translations
            output_path: Path for output file
            metadata: Optional document metadata
        
        Returns:
            Path to created document
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        self.doc = Document()
        
        # Set default style
        self._set_default_style()
        
        # Add metadata if provided
        if metadata:
            self._add_metadata(metadata)
        
        # Process each page
        for page in pages:
            await self._process_page(page)
        
        # Save document
        self.doc.save(str(output_path))
        
        self.logger.info(f"Saved DOCX to {output_path}")
        return output_path
    
    def _set_default_style(self):
        """Set default document styles."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        
        # Enable complex script for Russian/Chinese
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    def _add_metadata(self, metadata: Dict[str, Any]):
        """Add document metadata."""
        core_props = self.doc.core_properties
        
        if 'title' in metadata:
            core_props.title = metadata['title']
        if 'author' in metadata:
            core_props.author = metadata['author']
        if 'subject' in metadata:
            core_props.subject = metadata['subject']
    
    async def _process_page(self, page: PageContent):
        """Process a single page and add its elements."""
        # Sort elements by position
        elements = sorted(page.elements, key=lambda e: e.position)
        
        for element in elements:
            await self._add_element(element)
        
        # Add page break after each page except the last
        # (DOCX handles pagination automatically, but this helps maintain structure)
        # self.doc.add_page_break()
    
    async def _add_element(self, element: DocumentElement):
        """Add a single element to the document."""
        # Get translated content or fall back to original
        content = element.translated_content or element.content
        
        if element.element_type == ElementType.HEADER:
            self._add_header(content, element.metadata.get('font_size', 14))
        
        elif element.element_type == ElementType.PARAGRAPH:
            self._add_paragraph(content)
        
        elif element.element_type == ElementType.TABLE and element.table_data:
            self._add_table(element.table_data)
        
        elif element.element_type == ElementType.LIST_ITEM:
            self._add_list_item(content)
        
        elif element.element_type == ElementType.IMAGE:
            self._add_image_placeholder(element, content)
        
        elif element.is_ocr:
            # OCR content as regular paragraph with note
            self._add_ocr_paragraph(content)
        
        else:
            # Default: add as paragraph
            self._add_paragraph(content)
    
    def _add_header(self, text: str, font_size: float = 14):
        """Add a header with appropriate level."""
        # Determine header level based on font size
        if font_size >= 18:
            style = 'Heading 1'
            config = self.styles["header1"]
        elif font_size >= 14:
            style = 'Heading 2'
            config = self.styles["header2"]
        else:
            style = 'Heading 3'
            config = self.styles["header3"]
        
        try:
            paragraph = self.doc.add_heading(text, level=int(style[-1]))
        except ValueError:
            paragraph = self.doc.add_paragraph(text)
            paragraph.style = 'Heading 3'
        
        # Apply additional styling
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.bold = config["bold"]
        run.font.size = Pt(config["size"])
        
        paragraph.paragraph_format.space_after = Pt(config["space_after"])
    
    def _add_paragraph(self, text: str):
        """Add a normal paragraph."""
        paragraph = self.doc.add_paragraph(text)
        paragraph.style = 'Normal'
        
        # Configure paragraph formatting
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_after = Pt(self.styles["normal"]["space_after"])
        
        # Handle line breaks in text
        if '\n' in text:
            paragraph.clear()
            lines = text.split('\n')
            for i, line in enumerate(lines):
                if i > 0:
                    paragraph.add_run('\n')
                paragraph.add_run(line)
    
    def _add_ocr_paragraph(self, text: str):
        """Add OCR-processed text with visual indicator."""
        paragraph = self.doc.add_paragraph()
        paragraph.style = 'Normal'
        
        # Add note about OCR
        ocr_note = paragraph.add_run("[OCR] ")
        ocr_note.font.size = Pt(8)
        ocr_note.font.italic = True
        
        # Add text
        paragraph.add_run(text)
        
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    def _add_table(self, table_data: TableData):
        """Add a table with proper formatting."""
        rows = table_data.rows
        headers = table_data.headers
        
        if not rows:
            return
        
        # Determine table dimensions
        num_rows = len(rows)
        num_cols = max(len(row) for row in rows) if rows else 0
        
        if num_cols == 0:
            return
        
        # Add caption if present
        if table_data.caption:
            caption = self.doc.add_paragraph(table_data.caption)
            caption.style = 'Caption'
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create table
        docx_table = self.doc.add_table(rows=num_rows + (1 if headers else 0), cols=num_cols)
        docx_table.style = 'Table Grid'
        docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths (equal distribution)
        for col in docx_table.columns:
            col.width = Cm(15 / num_cols)  # Adjust based on page width
        
        # Fill headers
        if headers:
            header_row = docx_table.rows[0].cells
            for i, cell_text in enumerate(headers[:num_cols]):
                cell = header_row[i]
                cell.text = str(cell_text) if cell_text else ""
                
                # Style header cells
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
        
        # Fill data rows
        start_row = 1 if headers else 0
        for row_idx, row in enumerate(rows):
            docx_row = docx_table.rows[row_idx + start_row]
            for col_idx, cell_text in enumerate(row[:num_cols]):
                cell = docx_row.cells[col_idx]
                cell.text = str(cell_text) if cell_text else ""
                
                # Style data cells
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(self.styles["table_text"]["size"])
    
    def _add_list_item(self, text: str):
        """Add a list item."""
        paragraph = self.doc.add_paragraph(style='List Bullet')
        paragraph.add_run(text)
        paragraph.paragraph_format.space_after = Pt(6)
    
    def _add_image_placeholder(self, element: DocumentElement, caption: str):
        """Add an image placeholder with caption."""
        # Add placeholder text
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        placeholder = paragraph.add_run("[Изображение]")
        placeholder.font.size = Pt(10)
        placeholder.font.italic = True
        
        # Add caption if available
        if element.metadata.get('ocr_processed'):
            caption_para = self.doc.add_paragraph(caption[:200] + "..." if len(caption) > 200 else caption)
            caption_para.style = 'Caption'
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    async def add_custom_content(
        self,
        content_type: str,
        content: Any,
        **kwargs
    ):
        """
        Add custom content to the document.
        
        Args:
            content_type: Type of content ('text', 'table', 'image', etc.)
            content: Content to add
            **kwargs: Additional parameters
        """
        if content_type == 'text':
            self.doc.add_paragraph(content)
        elif content_type == 'heading':
            level = kwargs.get('level', 1)
            self.doc.add_heading(content, level=level)
        elif content_type == 'page_break':
            self.doc.add_page_break()
